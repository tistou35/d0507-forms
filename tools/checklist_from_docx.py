#!/usr/bin/env python3
# ============================================================
# checklist_from_docx.py — ถอดเช็กลิสต์จาก .docx เป็นรูปแบบของ
#                          collection "checklists" (เช็กลิสต์ที่กดติ๊กได้ในระบบ)
#
#   python3 tools/checklist_from_docx.py PCK            ดูผลก่อน
#   python3 tools/checklist_from_docx.py PCK --json     พิมพ์ JSON ออกมา
#
# ── ทำไมถอดด้วยเครื่อง ไม่พิมพ์เอง ────────────────────────────
# เอกสารควบคุมพิมพ์ผิดไม่ได้ ถอดจากไฟล์ต้นฉบับตรง ๆ แปลว่าข้อความในระบบ
# ตรงกับกระดาษทุกตัวอักษร และถ้าต้นฉบับแก้ ก็รันใหม่ได้ทั้งใบ
#
# ── รูปแบบที่ได้ ─────────────────────────────────────────────
#   items: [ {t: "1. DOCUMENTATION", sec: true},
#            {t: "Maintenance Release", a: "Valid maintenance release …", sub: []} ]
# แถวหัวข้อคือแถวที่ผสานช่องเป็นช่องเดียว (ไม่มีช่อง ☐ ต่อท้าย)
# ============================================================
import json
import os
import re
import sys

import docx
from docx.oxml.ns import qn

HERE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SRC_DIR = os.path.dirname(HERE)
REG = os.path.join(HERE, 'forms_register.json')


def cells_of(row):
    """ช่องที่ไม่ซ้ำในแถว — Word นับช่องที่ผสานไว้ซ้ำหลายครั้ง"""
    out = []
    for c in row.cells:
        if not out or out[-1]._tc is not c._tc:
            out.append(c)
    return out


def parse(path):
    d = docx.Document(path)
    # ตารางรายการคือตารางที่มีแถวเยอะที่สุด — ตารางอื่นเป็นหัวกระดาษและช่องลงนาม
    tbl = max(d.tables, key=lambda t: len(t.rows))
    items = []
    for row in tbl.rows:
        cs = cells_of(row)
        txt = cs[0].text.strip()
        if not txt:
            continue
        # แถวหัวข้อ: ช่องเดียวทั้งแถว และไม่มีช่องติ๊ก
        if len(cs) == 1:
            items.append({'t': re.sub(r'\s+', ' ', txt), 'sec': True})
            continue
        if txt.upper() in ('INSPECTION ITEM',):      # แถวหัวตาราง ไม่ใช่รายการ
            continue
        # ในช่องเดียวกันมีสองย่อหน้า: ป้ายกำกับ แล้วคำอธิบาย
        paras = [p.text.strip() for p in cs[0].paragraphs if p.text.strip()]
        label = re.sub(r'\s+', ' ', paras[0]) if paras else txt
        desc = re.sub(r'\s+', ' ', ' · '.join(paras[1:])) if len(paras) > 1 else ''
        items.append({'t': label, 'a': desc, 'sub': []})
    return items


def build(abbr):
    reg = json.load(open(REG, encoding='utf-8'))
    f = next((x for x in reg['forms'] if x['abbr'] == abbr), None)
    if not f:
        sys.exit('%s: ไม่มีในทะเบียน' % abbr)
    path = os.path.join(SRC_DIR, f.get('docx') or '')
    if not os.path.exists(path):
        sys.exit('%s: ไม่พบไฟล์ %s' % (abbr, f.get('docx')))
    items = parse(path)
    eff = f.get('eff') or ''
    # ทะเบียนเก็บ "18 JUL 2024" แต่ checklists เก็บ YYYY-MM-DD
    M = ['JAN', 'FEB', 'MAR', 'APR', 'MAY', 'JUN',
         'JUL', 'AUG', 'SEP', 'OCT', 'NOV', 'DEC']
    m = re.match(r'^(\d{1,2})\s+([A-Z]{3})\s+(\d{4})$', eff.strip().upper())
    iso = '%s-%02d-%02d' % (m.group(3), M.index(m.group(2)) + 1, int(m.group(1))) if m else ''
    return {
        'doc': f['doc'], 'code': f.get('code', ''),
        'issue': f.get('iss', '01'), 'rev': f.get('rev', '00'), 'eff': iso,
        'layout': 'one',       # คำอธิบายยาว สองคอลัมน์จะตัดคำกลางประโยค
        'title': {'th': f.get('th', ''), 'en': f.get('t', '')},
        'items': items,
    }


if __name__ == '__main__':
    args = [a for a in sys.argv[1:] if not a.startswith('--')]
    if not args:
        sys.exit('ใช้: python3 tools/checklist_from_docx.py PCK [--json]')
    out = build(args[0].upper())
    if '--json' in sys.argv:
        print(json.dumps(out, ensure_ascii=False))
        sys.exit(0)
    secs = [i for i in out['items'] if i.get('sec')]
    rows = [i for i in out['items'] if not i.get('sec')]
    print('%s · %s · ฉบับ %s/%s · %d หัวข้อ · %d รายการ'
          % (out['doc'], out['code'], out['issue'], out['rev'], len(secs), len(rows)))
    for i in out['items']:
        if i.get('sec'):
            print('\n  ▸ %s' % i['t'])
        else:
            print('     ☐ %-26s %s' % (i['t'][:26], i.get('a', '')[:66]))
