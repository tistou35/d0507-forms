#!/usr/bin/env python3
# ============================================================
# bump_code.py — เลื่อนอักษรท้ายเลขกำกับฟอร์มเมื่อออกฉบับแก้ไข
#
#   python3 tools/bump_code.py EFC          เลื่อนจริง (ทะเบียน + .docx)
#   python3 tools/bump_code.py EFC -n       ดูว่าจะเปลี่ยนเป็นอะไร ไม่เขียน
#
# ── กติกา ────────────────────────────────────────────────────
# เลขกำกับมีรูปแบบ  IM-EFC-303-A  โดยอักษรตัวท้ายบอก "รุ่นของแบบฟอร์ม"
# ทุกครั้งที่ ISSUE NO. หรือ REVISION NO. เปลี่ยน อักษรตัวท้ายต้องเลื่อนหนึ่งตัว
#   ISSUE 01 → 02        A → B
#   REVISION 00 → 01     A → B
# ไม่ว่าจะแก้แบบไหน อักษรก็เลื่อนเหมือนกัน
#
# ── ทำไมต้องมีเครื่องมือ ─────────────────────────────────────
# เลขนี้อยู่สามที่: หัวเอกสาร ท้ายเอกสาร และทะเบียน
# แก้มือแล้วลืมที่ใดที่หนึ่ง จะได้เอกสารที่อ้างเลขกำกับคนละตัวในใบเดียวกัน
# ซึ่งเป็นข้อบกพร่องที่ผู้ตรวจจับได้ทันทีและอธิบายยาก
# ============================================================
import json, os, re, sys

HERE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DOCX_DIR = os.path.dirname(HERE)
REG = os.path.join(HERE, 'forms_register.json')
CODE_RE = re.compile(r'^(.*-)([A-Z])$')


def next_letter(code):
    m = CODE_RE.match(code or '')
    if not m:
        raise ValueError('เลขกำกับ %r ไม่ลงท้ายด้วยอักษร A-Z — เลื่อนเองไม่ได้' % code)
    head, ch = m.group(1), m.group(2)
    if ch == 'Z':
        raise ValueError('เลขกำกับถึง Z แล้ว (%s) — ต้องตัดสินใจเรื่องรูปแบบใหม่' % code)
    return head + chr(ord(ch) + 1)


def main(argv):
    abbrs = [a for a in argv if not a.startswith('-')]
    dry = '-n' in argv
    if not abbrs:
        sys.exit('ใช้: bump_code.py <ABBR> [-n]')

    reg = json.load(open(REG, encoding='utf-8'))
    bad = 0
    for abbr in abbrs:
        f = next((x for x in reg['forms'] if x['abbr'] == abbr), None)
        if not f:
            print('🔴 %s ไม่มีในทะเบียน' % abbr); bad += 1; continue
        old = f.get('code') or ''
        try:
            new = next_letter(old)
        except ValueError as e:
            print('🔴 %s %s' % (abbr, e)); bad += 1; continue

        print('%s  %s → %s' % (abbr, old, new))

        src = os.path.join(DOCX_DIR, f.get('docx') or '')
        hits = 0
        if os.path.exists(src):
            import docx
            d = docx.Document(src)
            def fix(par):
                full = ''.join(r.text for r in par.runs)
                if old not in full or not par.runs:
                    return 0
                par.runs[0].text = full.replace(old, new)
                for r in par.runs[1:]:
                    r.text = ''
                return 1
            for p in d.paragraphs:
                hits += fix(p)
            for t in d.tables:
                for row in t.rows:
                    for c in row.cells:
                        for p in c.paragraphs:
                            hits += fix(p)
            print('   เอกสาร: พบและแก้ %d จุด' % hits)
            if not dry:
                d.save(src)
        else:
            print('   ⚠️ ไม่พบไฟล์ต้นฉบับ %s — แก้เฉพาะทะเบียน' % (f.get('docx') or '(ไม่ระบุ)'))

        if not hits and os.path.exists(src):
            print('   🔴 ไม่เจอเลขกำกับในเอกสารเลย — ตรวจว่าเลขในทะเบียนตรงกับที่พิมพ์จริงไหม')
            bad += 1
        f['code'] = new

    if dry:
        print('\nไม่ได้เขียนไฟล์ — เอา -n ออกเพื่อทำจริง')
        return 0 if not bad else 1
    json.dump(reg, open(REG, 'w', encoding='utf-8'), ensure_ascii=False, indent=1)
    open(REG, 'a', encoding='utf-8').write('\n')
    print('\nเขียนทะเบียนแล้ว — อย่าลืม build.py และสร้างแม่แบบใหม่')
    return 0 if not bad else 1


if __name__ == '__main__':
    sys.exit(main(sys.argv[1:]))
