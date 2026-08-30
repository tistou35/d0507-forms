#!/usr/bin/env python3
# ============================================================
# stamp_car.py — ใส่เลขกำกับลง D-0507-CAR-001 ให้ครบทุกที่
#
#   python3 tools/stamp_car.py            ดูก่อนว่าจะเขียนอะไร
#   python3 tools/stamp_car.py --write    เขียนจริง
#
# ── ที่มาของเลข ──────────────────────────────────────────────
# รูปแบบในทะเบียนคือ <แผนก>-<ตัวย่อ>-<เลขสามหลัก>-<ตัวอักษร>
# เลข 301 คือเลขแรกของเอกสารตระกูลใหม่ — ตรงกับ QA-DRF-301-A ·
# QA-EFM-301-B · QA-NCHR-301-A และอีก 13 ใบในแผนกอื่นที่ขึ้นต้นด้วย 301
# เลข 302 ขึ้นไปใช้เมื่อตระกูลนั้นมีใบแรกอยู่แล้ว (QA-DRF-302-A ของ DRC)
# CAR ยังไม่มีใบในตระกูลนี้เลย จึงเป็น QA-CAR-301-A ตัวอักษร A เพราะเป็น
# ฉบับพิมพ์แรก (ฉบับ 01 / แก้ไข 00)
#
# ── สามที่บนกระดาษ ──────────────────────────────────────────
#   1. ในเอกสาร   แถวใต้ชื่อฟอร์มในตารางหัวเรื่อง
#   2. หัวกระดาษ  Word header
#   3. ท้ายกระดาษ Word footer
# ทั้งสามที่ต้องตรงกัน ไม่งั้นกระดาษแผ่นเดียวอ้างเลขกำกับคนละเลข
# ซึ่งเป็นสิ่งที่ผู้ตรวจเห็นทันที (เคยเจอกับ MTC และ PWR มาแล้ว)
# ============================================================
import copy
import json
import os
import sys

import docx
from docx.oxml.ns import qn

HERE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SRC_DIR = os.path.dirname(HERE)
REG = os.path.join(HERE, 'forms_register.json')

ABBR = 'CAR'
CODE = 'QA-CAR-301-A'


def reg_row():
    r = json.load(open(REG, encoding='utf-8'))
    return r, next(x for x in r['forms'] if x['abbr'] == ABBR)


def retext(par, old, new):
    """แทนคำทีละ run — ไม่รวม run เข้าด้วยกัน ตัวหนา/ขนาดจึงไม่เพี้ยน"""
    for run in par.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return True
    # ข้อความถูกหั่นข้าม run — เขียนทั้งย่อหน้าลง run แรกแทน
    if old in par.text:
        full = par.text.replace(old, new)
        for run in par.runs[1:]:
            run.text = ''
        par.runs[0].text = full
        return True
    return False


def main():
    write = '--write' in sys.argv
    reg, f = reg_row()
    path = os.path.join(SRC_DIR, f['docx'])
    eff = f.get('eff', '')
    d = docx.Document(path)
    did = []

    # ── 1. ในเอกสาร — แถวใหม่ใต้ชื่อฟอร์ม ──
    t = d.tables[0]
    line = '%s    EFF: %s    ·    Compliance Monitoring (QA)  ·  D-0507 ATO' % (CODE, eff)
    if CODE not in t.rows[-1].cells[0].text:
        proto = t.rows[-1]._tr
        new = copy.deepcopy(proto)
        for tc in new.findall(qn('w:tc')):
            ps = tc.findall(qn('w:p'))
            for p in ps[1:]:
                tc.remove(p)
            runs = ps[0].findall(qn('w:r'))
            for r in runs[1:]:
                ps[0].remove(r)
            if runs:
                for tnode in runs[0].findall(qn('w:t')):
                    runs[0].remove(tnode)
                el = runs[0].makeelement(qn('w:t'), {})
                el.set(qn('xml:space'), 'preserve')
                el.text = line
                runs[0].append(el)
        proto.addnext(new)
        did.append('ในเอกสาร: เพิ่มแถว "%s"' % line)

    # ── 2. หัวกระดาษ ──
    for p in d.sections[0].header.paragraphs:
        if 'Corrective Action Request' in p.text and CODE not in p.text:
            if retext(p, 'Corrective Action Request (CAR)',
                      'Corrective Action Request (CAR)  —  ' + CODE):
                did.append('หัวกระดาษ: %s' % p.text.strip())

    # ── 3. ท้ายกระดาษ ──
    for p in d.sections[0].footer.paragraphs:
        if 'D-0507-CAR-001' in p.text and CODE not in p.text:
            if retext(p, 'D-0507-CAR-001  |  Issue 01 Rev 00',
                      'D-0507-CAR-001  |  %s  |  Issue 01 Rev 00  |  EFF: %s' % (CODE, eff)):
                did.append('ท้ายกระดาษ: %s' % p.text.strip())

    # ── 4. ทะเบียน ──
    if f.get('code') != CODE:
        did.append('ทะเบียน: code ว่าง → %s' % CODE)

    if not did:
        print('มีเลขกำกับครบทุกที่แล้ว — ไม่ต้องทำอะไร')
        return
    for x in did:
        print('  ·', x)
    if not write:
        print('\n(ดูอย่างเดียว — ใส่ --write เพื่อเขียนจริง)')
        return
    d.save(path)
    f['code'] = CODE
    json.dump(reg, open(REG, 'w', encoding='utf-8'), ensure_ascii=False, indent=1)
    print('\nเขียนแล้ว: %s และ forms_register.json' % f['docx'])


if __name__ == '__main__':
    main()
