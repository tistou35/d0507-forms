#!/usr/bin/env python3
# ============================================================
# model_chk.py — ทำเช็กลิสต์ NPC / EPC ให้ตรงกับเครื่องยนต์ของแต่ละรุ่น
#
#   python3 tools/model_chk.py            ทำทั้ง NPC และ EPC ทั้งสามรุ่น
#   python3 tools/model_chk.py NPC
#
# ── ปัญหาที่แก้ ──────────────────────────────────────────────
# ฉบับที่ใช้อยู่มาจากไฟล์เดียวกันทั้งสามรุ่น จึงมีทั้งรายการของเครื่องยนต์
# คาร์บูเรเตอร์และของหัวฉีดปนกันอยู่ในใบเดียว
#
#   C172M / C172N   Lycoming O-320 คาร์บูเรเตอร์ · น้ำมันไหลด้วยแรงโน้มถ่วง
#                   ไม่มีปั๊มน้ำมันไฟฟ้า · มี carb heat และ primer
#   C172S           Lycoming IO-360 หัวฉีด · มีปั๊มน้ำมันไฟฟ้า (AUX)
#                   ไม่มี carb heat และไม่มี primer
#
# ใบของ M/N จึงสั่งให้เปิดปั๊มที่เครื่องบินไม่มี ส่วนใบของ S สั่งให้ดึง carb heat
# ที่ไม่มีคันโยก — ทั้งสองอย่างคือรายการที่ทำตามไม่ได้ในห้องนักบิน
#
# ── หลักในการแก้ ────────────────────────────────────────────
# ตัดเฉพาะรายการของอุปกรณ์ที่เครื่องบินรุ่นนั้นไม่มี และแก้ค่าที่ POH ต่างกัน
# ไม่แต่งขั้นตอนใหม่ขึ้นเอง — ขั้นตอนที่ยังขาด (เช่นการไล่น้ำมันก่อนสตาร์ต
# ของเครื่องหัวฉีด) ปล่อยไว้ให้ครูการบินเติมจาก POH ของเครื่องจริง
# ทุกตัวเลขในตารางด้านล่างต้องถูกทวนกับ POH ของแต่ละลำก่อนอนุมัติ
#
# ── วิธี ─────────────────────────────────────────────────────
# เริ่มจากฉบับคอลัมน์เดียว (*-1col-archive.docx) เพราะการตัดแถวทำได้ตรงไปตรงมา
# แก้เสร็จแล้วค่อยจัดเป็นหลายคอลัมน์ด้วยโค้ดชุดเดิมใน build_2col_chk.py
# ตัดแถวจากฉบับหลายคอลัมน์โดยตรงไม่ได้ เพราะหนึ่งแถวถือรายการของสองสามคอลัมน์
# ที่ไม่เกี่ยวกันเลย ตัดแถวหนึ่งคือตัดรายการของหมวดอื่นทิ้งไปด้วย
# ============================================================
import copy
import os
import sys

import docx
from docx.oxml.ns import qn

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from build_2col_chk import (blank_tc, cells_from, empty_pair, flat, rows_of,  # noqa: E402
                            split, widths, W_GAP)

HERE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SRC_DIR = os.path.dirname(HERE)

COLS = {'NPC': 3, 'EPC': 2}          # เท่าฉบับที่ใช้อยู่ อย่าให้จำนวนหน้าเปลี่ยน
SEQ = {'M': 303, 'N': 304, 'S': 305}
PREFIX = {'NPC': 'NP-CHK', 'EPC': 'EP-CHK'}
EFF = '25 AUG 2026'

# ── สิ่งที่รุ่นนั้นไม่มี — ตัดทั้งแถว ─────────────────────────
# เทียบจากข้อความในช่องซ้าย (ไม่สนตัวพิมพ์เล็กใหญ่)
DROP = {
    'NPC': {
        # M/N ไม่มีปั๊มน้ำมันไฟฟ้า น้ำมันไหลลงจากปีกด้วยแรงโน้มถ่วง
        'M': ['fuel pump'],
        'N': ['fuel pump'],
        # S เป็นเครื่องหัวฉีด ไม่มีคันโยก carb heat และไม่มี primer
        'S': ['carb heat', 'primer'],
    },
    'EPC': {
        'M': ['fuel pump'],
        'N': ['fuel pump'],
        'S': ['carb heat', 'primer'],
    },
}

# ── ค่าที่ POH ต่างกันจริง ───────────────────────────────────
# (ข้อความเดิม, ข้อความใหม่) — แทนทีละ run ให้ตัวหนา/ขนาดไม่เพี้ยน
RETEXT = {
    'NPC': {
        # runup ของ O-320 อยู่ที่ 1,700 RPM ส่วน 1,800 เป็นค่าของ IO-360
        'M': [('Throttle to 1,800 RPM', 'Throttle to 1,700 RPM')],
        'N': [('Throttle to 1,800 RPM', 'Throttle to 1,700 RPM')],
        'S': [
            ('max 125 RPM drop', 'max 150 RPM drop'),
            # 172S ตั้งแต่รุ่นปี 1996 จำกัดแฟลปที่ 30 องศา ไม่มี 40
            ('FULL (40°) as required', 'FULL (30°) as required'),
            ('40° (full)', '30° (full)'),
            ('40°', '30°'),
        ],
    },
    'EPC': {
        'M': [], 'N': [],
        # best glide ของ 172S ที่น้ำหนักสูงสุดคือ 68 KIAS
        'S': [('65 KIAS (best glide)', '68 KIAS (best glide)'),
              ('Best glide — 65 KIAS', 'Best glide — 68 KIAS')],
    },
}

# ── แก้ชื่อหัวข้อ ตามที่ผู้ควบคุมเอกสารสั่ง ───────────────────
# ชื่อเดิมอ่านเหมือนเป็นหมวด "ก่อนวิ่งขึ้น" ทั้งที่เนื้อในคือขั้นตอน runup
HEADINGS = [('BEFORE TAKEOFF (RUNUP)', 'RUNUP [Completed Before Takeoff]')]


def item_text(tr):
    tcs = tr._tr.findall(qn('w:tc'))
    return ''.join(tcs[0].itertext()) if tcs else ''


def retext_tr(tr, pairs):
    """แทนคำในทุก run ของแถว — ไม่รวม run เข้าด้วยกัน รูปแบบตัวอักษรจึงคงเดิม"""
    n = 0
    for t in tr._tr.iter(qn('w:t')):
        for a, b in pairs:
            if a in (t.text or ''):
                t.text = t.text.replace(a, b)
                n += 1
    return n


def apply_edits(tbl, abbr, model):
    """ตัดแถวที่รุ่นนี้ไม่มี แล้วแก้ค่าที่ต่างกัน คืนสรุปว่าทำอะไรไปบ้าง"""
    drop_keys = DROP[abbr][model]
    dropped, changed = [], 0
    for tr in list(tbl.rows):
        tcs = tr._tr.findall(qn('w:tc'))
        is_sec = len(tcs) == 1 or tr.cells[0]._tc is tr.cells[1]._tc
        if is_sec:
            changed += retext_tr(tr, HEADINGS)
            continue
        txt = item_text(tr)
        if any(k in txt.lower() for k in drop_keys):
            act = ''.join(tcs[1].itertext()) if len(tcs) > 1 else ''
            dropped.append('%s — %s' % (txt.replace('☐', '').strip(), act.strip()))
            tr._tr.getparent().remove(tr._tr)
            continue
        changed += retext_tr(tr, RETEXT[abbr][model])
    return dropped, changed


def relayout(d, cols):
    """จัดตารางคอลัมน์เดียวให้เป็นหลายคอลัมน์ — ใช้โค้ดชุดเดียวกับ build_2col_chk"""
    W_ITEM, W_ACT = widths(cols)
    tbl = d.tables[-1]
    parts = [flat(g) for g in split(rows_of(tbl), cols)]
    if any(not p for p in parts):
        sys.exit('แบ่ง %d คอลัมน์ไม่ได้' % cols)

    new = copy.deepcopy(tbl._tbl)
    for tr in new.findall(qn('w:tr')):
        new.remove(tr)
    grid = new.find(qn('w:tblGrid'))
    for gc in grid.findall(qn('w:gridCol')):
        grid.remove(gc)
    from docx.oxml import OxmlElement
    for n in range(cols):
        if n:
            gc = OxmlElement('w:gridCol'); gc.set(qn('w:w'), str(W_GAP)); grid.append(gc)
        for w in (W_ITEM, W_ACT):
            gc = OxmlElement('w:gridCol'); gc.set(qn('w:w'), str(w)); grid.append(gc)

    for i in range(max(len(p) for p in parts)):
        tr = OxmlElement('w:tr')
        for n, part in enumerate(parts):
            if n:
                tr.append(blank_tc(W_GAP))
            src = cells_from(part[i][0], part[i][1], W_ITEM, W_ACT) if i < len(part) \
                else empty_pair(W_ITEM, W_ACT)
            for tc in src:
                tr.append(tc)
        new.append(tr)

    tbl._tbl.addnext(new)
    tbl._tbl.getparent().remove(tbl._tbl)
    return len(new.findall(qn('w:tr')))


def stamp(d, abbr, model):
    """เลขกำกับ วันมีผล และชื่อรุ่น บนหัวและท้ายกระดาษ

    เนื้อหาเปลี่ยน = ฉบับพิมพ์เปลี่ยน ตัวอักษรท้ายเลขกำกับจึงต้องเลื่อน -A → -B
    ตามกฎใน CLAUDE.md และต้องเลื่อนทั้งหัวและท้าย ไม่งั้นกระดาษแผ่นเดียว
    อ้างเลขกำกับสองเลข ซึ่งผู้ตรวจเห็นทันที
    """
    import re
    new_code = '%s-%d-B' % (PREFIX[abbr], SEQ[model])
    pat = re.compile(re.escape(PREFIX[abbr]) + r'-\d{3}-[A-Z]')
    hits = 0

    for p in d.paragraphs:
        for r in p.runs:
            t = pat.sub(new_code, r.text)
            t = t.replace('EFF: 11 NOV 2024', 'EFF: ' + EFF)
            t = t.replace('Cessna 172 Series', 'Cessna 172%s' % model)
            t = t.replace('Cessna 172  |', 'Cessna 172%s  |' % model)
            # คำโปรยเตือนอ้าง POH ของรุ่นด้วย ไม่ใช่ "Cessna 172" ลอย ๆ
            t = t.replace('Cessna 172 POH', 'Cessna 172%s POH' % model)
            if t != r.text:
                r.text = t
                hits += 1
    hits += tab_seps(d)
    return hits, new_code


def tab_seps(d):
    """เติมตัวคั่นให้ข้อความที่คั่นด้วยแท็บ

    ใน Word แท็บดันข้อความไปชิดขวาตาม tab stop แต่ตัวแปลงเป็น PDF ทิ้งแท็บนั้น
    หัวกระดาษจึงพิมพ์ออกมาเป็น "…Co., Ltd.NP-CHK-303-B" ติดกันเป็นคำเดียว
    ทุกแผ่น เจอแบบเดียวกันมาแล้วที่ท้ายกระดาษของ ALR

    แท็บมักอยู่คนละ run กับข้อความข้างหน้า จึงต้องดูทั้งย่อหน้า ไม่ใช่ทีละ run
    """
    n = 0
    for p in d.paragraphs:
        if '\t' not in p.text or '\t·' in p.text:
            continue
        for r in p.runs:
            if r.text.startswith('\t') and len(r.text) > 1:
                r.text = '\t·  ' + r.text[1:]
                n += 1
                break
            if r.text == '\t':
                r.text = '\t·  '
                n += 1
                break
    return n


def build(abbr, model):
    arch = os.path.join(SRC_DIR, 'D-0507-%s-001-1col-archive.docx' % abbr)
    if not os.path.exists(arch):
        sys.exit('ไม่พบ %s' % arch)
    out = os.path.join(SRC_DIR, 'D-0507-%s-%s-001.docx' % (abbr, model))

    d = docx.Document(arch)
    dropped, changed = apply_edits(d.tables[-1], abbr, model)
    rows = relayout(d, COLS[abbr])
    hits, code = stamp(d, abbr, model)
    d.save(out)

    print('  C172%s  %-14s %2d คอลัมน์ × %2d แถว · ตัด %d รายการ · แก้ค่า %d จุด · หัว/ท้าย %d'
          % (model, code, COLS[abbr], rows, len(dropped), changed, hits))
    for x in dropped:
        print('      ตัดออก: %s' % x)
    return out


if __name__ == '__main__':
    want = [a.upper() for a in sys.argv[1:]] or ['NPC', 'EPC']
    for abbr in want:
        if abbr not in COLS:
            sys.exit('ไม่รู้จัก %s — มีแค่ NPC และ EPC' % abbr)
        print('%s:' % abbr)
        for m in ('M', 'N', 'S'):
            build(abbr, m)
