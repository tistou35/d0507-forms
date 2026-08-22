# -*- coding: utf-8 -*-
"""D-0507-VSR-001 ฉบับแก้ไข 01 — เก็บสิ่งที่กระดาษถามอยู่แล้วให้ครบ

   cd "<โฟลเดอร์ Manual revision>" && python3 d0507-forms/tools/build_vsr.py

ฉบับปี 2020 มีสามหัวข้อในส่วนบรรยาย (เหตุการณ์ · ปัจจัยที่เกี่ยวข้อง · ข้อเสนอแนะ)
แต่ทั้งสามเป็นหัวข้อเปล่าไม่มีป้ายคู่ช่อง ตัวจับคู่ token จึงวางไม่ได้ ผลคือคำตอบ
สองในสามข้อไม่เคยขึ้น PDF และ "ปัจจัยที่เกี่ยวข้อง" ไม่มีช่องกรอกในระบบด้วยซ้ำ —
กระดาษถาม แต่ไม่มีใครตอบได้

⚠️ ห้ามแก้ไฟล์นี้ด้วย regex — ต้นฉบับมี <w:tr> ซ้อนใน <w:tr> อยู่ห้าจุด
   (XML ถูกต้องแต่ผิดหลัก OOXML) การตัดแถวด้วย regex จะได้ขอบเขตผิด
   แล้วเขียนทับผิดแถวโดยไม่มีอะไรฟ้อง ลองมาแล้วสองรอบ ใช้ python-docx เท่านั้น

⚠️ รันได้ครั้งเดียวจากฉบับ Rev 00 · จะไม่เขียนทับ archive ที่มีอยู่แล้ว
"""
import copy, os, shutil
import docx

SRC = OUT = 'D-0507-VSR-001.docx'
ARCHIVE = 'D-0507-VSR-001-rev00-archive.docx'
W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'


def texts(el):
    return el.findall('.//' + W + 't')


def set_text(el, s):
    """เขียนข้อความลง <w:t> ตัวแรกของ element แล้วล้างที่เหลือ

    ไม่ลบ run ทิ้งเพราะ run พก rPr (ฟอนต์ ขนาด สี) ไว้ ลบแล้วต้องสร้างใหม่ให้เหมือนเดิม
    """
    ts = texts(el)
    if not ts:
        raise SystemExit('element นี้ไม่มี <w:t> ให้เขียน')
    ts[0].text = s
    ts[0].set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    for t in ts[1:]:
        t.text = ''


def unwrap_rows(d):
    """แกะ <w:tr> ที่ถูกห่อด้วย <w:tr> อีกชั้นออก

    ต้นฉบับมีแถวหัวข้อห้าแถวที่เขียนเป็น <w:tr><w:tr>...</w:tr></w:tr> แถวนอกไม่มี
    <w:tc> เลย ผิดหลัก OOXML — Word เปิดได้ แต่ตัวแปลงของ Drive ทิ้งทั้งคู่
    ผลคือหัวข้อ DESCRIPTION / CONTRIBUTING FACTORS / SUGGESTED ACTION
    ไม่เคยขึ้นใน PDF สักใบตั้งแต่ออกเอกสารมา เพิ่งเห็นตอนเปิด PDF เทียบกับ .docx
    """
    n = 0
    for tbl in d.tables:
        for tr in list(tbl._tbl.findall(W + 'tr')):
            inner = tr.findall(W + 'tr')
            if len(inner) == 1 and tr.find(W + 'tc') is None:
                tbl._tbl.replace(tr, inner[0])
                n += 1
    return n


def main():
    if os.path.exists(ARCHIVE):
        raise SystemExit('มี %s อยู่แล้ว — ย้ายออกก่อนถ้าตั้งใจสร้างใหม่' % ARCHIVE)
    d = docx.Document(SRC)
    if len(d.tables) != 7:
        raise SystemExit('คาดว่ามี 7 ตาราง พบ %d — โครงเอกสารเปลี่ยนไป หยุดก่อน' % len(d.tables))
    t4, t5 = d.tables[4], d.tables[5]
    if len(t4.rows) != 7 or len(t5.rows) != 8:
        raise SystemExit('จำนวนแถวไม่ตรงที่คาด — หยุดก่อน')
    shutil.copy(SRC, ARCHIVE)

    # ── ซ่อมแถวหัวข้อที่ถูกห่อซ้อนก่อน ไม่งั้นหัวข้อไม่ขึ้นใน PDF ──
    print('   แกะแถวที่ห่อซ้อน %d แถว' % unwrap_rows(d))

    # ── ชนิดรายงาน: "Other:" ต้องมีที่รับคำตอบ ──
    set_text(t4.rows[6]._tr,
             '☐ Safety Hazard     ☐ Incident     ☐ Near-Miss     '
             '☐ Unsafe Act / Condition     ☐ Other:  {{evOther}}')

    # ── ส่วนบรรยาย: วาง token ลงช่องว่างใต้แต่ละหัวข้อ ──
    for ri, tok in ((2, '{{descr}}'), (5, '{{contributing}}'), (7, '{{suggest}}')):
        set_text(t5.rows[ri]._tr, tok)

    # ── หัวข้อไฟล์แนบ — ยืมโครงแถวหัวข้อและแถวช่องว่างที่มีอยู่ ──
    head = copy.deepcopy(t5.rows[6]._tr)
    set_text(head, 'ATTACHMENTS (photographs submitted with this report)')
    blank = copy.deepcopy(t5.rows[3]._tr)
    set_text(blank, '{{photo1}}     {{photo2}}')
    t5._tbl.append(head)
    t5._tbl.append(blank)

    # ── เลขกำกับและวันมีผล — อยู่ทั้งหัวตารางแรกและบรรทัดท้ายเอกสาร ──
    n = 0
    for t in d.element.body.iter(W + 't'):
        if t.text and 'SM-VSR-303-A' in t.text:
            t.text = t.text.replace('SM-VSR-303-A', 'SM-VSR-303-B')
            n += 1
        if t.text and 'EFF: 01 MAR 2020' in t.text:
            t.text = t.text.replace('EFF: 01 MAR 2020', 'EFF: 23 AUG 2026')
            n += 1
    if n < 2:
        raise SystemExit('แทนเลขกำกับได้แค่ %d จุด — คาดว่าอย่างน้อย 2' % n)

    d.save(OUT)
    print('เขียน %s · เก็บฉบับเดิมไว้ที่ %s · แทนเลขกำกับ %d จุด' % (OUT, ARCHIVE, n))


main()
