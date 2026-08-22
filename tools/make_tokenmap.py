#!/usr/bin/env python3
# ============================================================
# make_tokenmap.py — สร้างแผนที่ token จากนิยามฟอร์ม + .docx ต้นฉบับ
#
#   python3 tools/make_tokenmap.py            # ทุกใบที่มีนิยามฟอร์ม
#   python3 tools/make_tokenmap.py VSR HIF
#   python3 tools/make_tokenmap.py -v VSR     # บอกเหตุผลที่จับคู่ไม่ได้
#
# ผลลัพธ์: gas/TokenMap.gs  — ข้อมูลที่ ImportTemplate.gs ใช้เติม token
# ลงใน Google Doc ที่แปลงมาจาก .docx โดยตรง เลย์เอาต์จึงตรงต้นฉบับ 100%
#
# แผนที่มีสามส่วน เพราะกระดาษวางช่องกรอกไว้สามแบบ
#   byLabel  ป้ายอยู่ในเซลล์ ช่องกรอกอยู่เซลล์ถัดไป -> "Name" | ____
#   byLine   ป้ายกับเส้นประอยู่บรรทัดเดียวกันในเซลล์ -> "Signature: ______"
#   boxes    ช่องติ๊ก ☐ เรียงตามลำดับที่ปรากฏ
#
# ที่จับคู่ไม่ได้จะถูกส่งต่อไปให้ ImportTemplate.gs พิมพ์ท้ายเอกสาร
# ให้คนวางมือ — ดีกว่าเดาแล้ววางผิดช่องโดยไม่มีใครรู้
# ============================================================
import json, os, re, sys, glob, unicodedata

HERE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DOCX = os.path.dirname(HERE)          # โฟลเดอร์ Manual revision

_ovp = os.path.join(HERE, 'formdefs', '_TOKENMAP_OVERRIDES.json')
OVERRIDES = json.load(open(_ovp, encoding='utf-8')) if os.path.exists(_ovp) else {}

# คำพ้องความหมาย — กระดาษกับนิยามฟอร์มเรียกของเดียวกันคนละชื่อ
SYN = {
    'signature': ['sign', 'signed', 'ลายเซ็น', 'ลงชื่อ'],
    'name':      ['ชื่อ', 'print'],
    'date':      ['วันที่', 'dated'],
    'comment':   ['remark', 'ความเห็น', 'หมายเหตุ'],
    'other':     ['อื่น', 'specify', 'ระบุ'],
}
STOP = {'the', 'of', 'a', 'to', 'if', 'no', 'or', 'and', 'optional', 'please', 'print'}


def norm(s):
    """ตัดเครื่องหมายและช่องว่างออก เทียบป้ายในเอกสารกับป้ายในนิยามฟอร์ม"""
    s = unicodedata.normalize('NFKC', str(s or '')).lower()
    s = re.sub(r'\(.*?\)', ' ', s)                 # ตัดวงเล็บอธิบาย
    s = re.sub(r'[^a-z0-9฀-๿]+', ' ', s)
    return s.strip()


def norm_full(s):
    """เหมือน norm แต่เก็บข้อความในวงเล็บไว้

    PWR มีทั้ง "ID / Passport No." และ "ID / Passport No. (Guardian)"
    ถ้าตัดวงเล็บทิ้งสองอันนี้จะกลายเป็นอันเดียวกัน แล้ววาง token ผิดแถว
    """
    s = unicodedata.normalize('NFKC', str(s or '')).lower()
    s = re.sub(r'[^a-z0-9฀-๿]+', ' ', s)
    return s.strip()


def words(s):
    """ชุดคำสำหรับให้คะแนนความใกล้เคียง รวมคำพ้องเข้าไปด้วย"""
    w = set(x for x in norm(s).split() if x and x not in STOP)
    for key, alts in SYN.items():
        if w & ({key} | set(alts)):
            w |= {key}
    return w


# คำที่กว้างเกินกว่าจะชี้ชัดได้ด้วยตัวเอง — "Name" เฉย ๆ ตรงกับได้หลายช่อง
GENERIC = {'name', 'date', 'time', 'signature', 'no', 'other'}


def score(a, b):
    """สัดส่วนคำที่ทับกัน — 1.0 คือเหมือนกันทุกคำ

    b คือป้ายบนกระดาษ มักสั้นกว่าป้ายในนิยามฟอร์ม ("Other" กับ "Other — specify")
    ถ้าคำของ b อยู่ใน a ครบและไม่ได้กว้างจนชี้ชัดไม่ได้ ให้ถือว่าใกล้เคียงมาก
    """
    wa, wb = words(a), words(b)
    if not wa or not wb: return 0.0
    base = len(wa & wb) / max(len(wa), len(wb))
    if wb <= wa and not wb <= GENERIC:
        return max(base, 0.8)
    return base


def cell_head(cell):
    """เซลล์ที่มีทั้งป้ายและเส้นประให้กรอกอยู่ในตัวเอง — คืนข้อความที่ต้องเก็บไว้

    "Aircraft Type: ______"        -> "Aircraft Type: "        (ตัดเส้นประทิ้ง)
    "Date:  ___ / ___ / ______"    -> "Date:  "                (โครงวันที่ทิ้งทั้งชุด)
    "Aircraft Reg.:  HS-___"       -> "Aircraft Reg.:  HS-"    (เก็บ HS- ที่พิมพ์ไว้)
    ไม่เข้าเงื่อนไข -> None แปลว่าเป็นป้ายเฉย ๆ ค่าไปอยู่เซลล์ถัดไป
    """
    if '_' not in cell or '\n' in cell.strip(): return None
    # ตัดหางที่เป็นแต่ตัวยึดตำแหน่ง (เส้นประ ทับ จุด เว้นวรรค) ออกให้หมด
    head = re.sub(r'[_\s/.\-]*$', '', cell)
    if not re.search(r'_', cell[len(head):]): return None
    # เหลือแต่ป้ายล้วน ๆ ยังไม่พอ ต้องมีตัวคั่นหรือข้อความนำหน้าจริง
    return head + (' ' if head and not head.endswith((' ', '-', ':')) else '')


def box_same(tok_label, doc_label):
    """ชื่อตัวเลือกของ token กับข้อความข้าง ☐ ในกระดาษ หมายถึงอันเดียวกันไหม

    ใช้การ "อยู่ใน" เป็นหลัก ไม่ใช่คะแนนคำทับกัน เพราะตัวเลือกจริงมักเป็นคำสั้น
    อย่าง No / Yes / High ซึ่งถูกตัดทิ้งเป็น stopword ไปหมดถ้าใช้ words()
    ข้อความข้าง ☐ ในกระดาษมักมีหางต่อท้าย ("No   Date notified to HT: ___")
    """
    a, b = norm_full(tok_label), norm_full(doc_label)
    if not a or not b: return True          # ไม่มีข้อมูลให้เทียบ อย่าเพิ่งเตือน
    if a == b or a in b or b in a: return True
    return score(tok_label, doc_label) >= 0.34


def docx_box_labels(path):
    """ป้ายกำกับของ ☐ แต่ละช่อง เรียงตามที่ปรากฏจริงในเอกสาร

    ใช้ตรวจว่า token ที่จะวางเรียงตรงกับกระดาษไหม — การนับจำนวนเท่ากันพิสูจน์
    ไม่ได้ว่าเรียงถูก ASF เคยจำนวนตรง (16 = 16) แต่เลื่อนไปหนึ่งช่องตั้งแต่ข้อที่ 5
    ทำให้ใบที่ออกมาติ๊กผิดข้อโดยไม่มีอะไรฟ้อง
    """
    import docx
    d = docx.Document(path)
    out = []
    for t in d.tables:
        for r in t.rows:
            seen = []
            for c in r.cells:
                if c._element in seen: continue
                seen.append(c._element)
                n = c.text.count('☐')
                if not n: continue
                # ข้อความที่ตามหลัง ☐ แต่ละตัวในเซลล์นี้ = ชื่อตัวเลือก
                # ("☐ Ground Course ☐ Flight Training" -> สองตัวเลือก)
                segs = [s.strip(' :·—-') for s in c.text.split('☐')[1:]]
                segs = [re.split(r'\n', s)[0].strip() for s in segs]
                if any(segs):
                    out += [(s or '') for s in (segs + [''] * n)[:n]]
                else:
                    # ☐ อยู่ในคอลัมน์ของตัวเอง — ชื่อรายการอยู่เซลล์แรกของแถว
                    lab = next((x.text.strip() for x in r.cells
                                if x.text.strip() and '☐' not in x.text), '')
                    out += [lab] * n
    return out


def docx_parts(path):
    """คืน (เซลล์ทั้งหมดตามลำดับ, บรรทัดที่มีเส้นประให้กรอก)"""
    import docx
    d = docx.Document(path)
    cells, lines = [], []
    for t in d.tables:
        for r in t.rows:
            seen = []
            for c in r.cells:
                if c._element in seen: continue
                seen.append(c._element)
                txt = c.text.strip()
                cells.append(txt)
                # "Signature: _______" / "☐ Other: ______" — ป้ายกับเส้นประอยู่ด้วยกัน
                # ไม่บังคับว่าต้องเต็มบรรทัด เพราะบางใบวางไว้กลางประโยครวมกับช่องติ๊ก
                for ln in txt.split('\n'):
                    for m in re.finditer(r'([^\n:：_☐]{2,40}?)\s*[:：]\s*(_{3,})', ln):
                        lines.append((m.group(1).strip(), m.group(2)))
                    m = re.match(r'^\s*([^_:：]{2,60}?)\s*(_{4,})\s*$', ln)   # ไม่มีโคลอน
                    if m: lines.append((m.group(1).strip(), m.group(2)))
    return cells, lines


def field_anchor(f):
    """ข้อความที่ใช้ค้นหาในเอกสาร — ชนิดช่องบอกใบ้ได้ว่าป้ายในกระดาษน่าจะเขียนว่าอะไร"""
    lab = (f.get('label') or {}).get('en') or ''
    if f.get('type') == 'sign' and 'sign' not in lab.lower():
        lab += ' signature'
    return lab


def build(abbr, verbose=False):
    fp = os.path.join(HERE, 'formdefs', abbr + '.json')
    dp = os.path.join(DOCX, 'D-0507-%s-001.docx' % abbr)
    if not os.path.exists(fp): return None, ['ไม่มีนิยามฟอร์ม'], None
    if not os.path.exists(dp): return None, ['ไม่มี .docx ต้นฉบับ'], None

    d = json.load(open(fp, encoding='utf-8'))
    # ส่วนอนุมัติเป็นของที่ระบบเพิ่มเข้ามา กระดาษเดิมส่วนใหญ่ไม่มี
    # จับคู่ไม่ได้เป็นเรื่องปกติ ต้องต่อท้ายเอกสารให้เอง ไม่ใช่ค้างไว้ให้วางมือ
    apParty = (d.get('route') or [{}, {}])[1].get('party') if len(d.get('route') or []) > 1 else None
    fields = []
    for s in d.get('sections', []):
        for f in s.get('fields', []):
            f = dict(f, _ap=(s.get('party') == apParty and apParty))
            fields.append(f)
    cells, lines = docx_parts(dp)
    short = [c for c in cells if c and len(c) < 90]
    # เก็บสองชั้น: ชั้นเต็มไว้แยกป้ายที่ต่างกันแค่วงเล็บ ชั้นย่อไว้จับคู่หลวม ๆ
    cellfull, cellset = {}, {}
    for c in short: cellfull.setdefault(norm_full(c), c)
    for c in short: cellset.setdefault(norm(c), c)
    # ป้ายเดียวกันโผล่ได้หลายที่ ("Phone Number" ของผู้โดยสารและของผู้ปกครอง)
    # จับคู่ได้เท่าจำนวนที่มีจริง ฝั่ง Apps Script จะไล่หาอันถัดไปเอง
    avail = {}
    for c in short: avail[c] = avail.get(c, 0) + 1

    by_label, by_line, by_cell, boxes, unmatched, approval, tables = [], [], [], [], [], [], []
    used_lines = set()

    def take(c):
        """จองเซลล์นี้ไว้หนึ่งครั้ง คืน True ถ้ายังเหลือให้จอง"""
        if avail.get(c, 0) <= 0: return False
        avail[c] -= 1
        return True

    for f in fields:
        if f.get('type') == 'static':
            continue
        tok = '{{sig_%s}}' % f['k'] if f.get('type') == 'sign' else '{{%s}}' % f['k']

        # ช่องติ๊กและตัวเลือก — เก็บไว้ใส่แทน ☐ ตามลำดับ
        if f.get('type') == 'check':
            lb = (f.get('label') or {}).get('en') or ''
            boxes.append({'tok': '{{k_%s}}' % f['k'], 'label': lb, 'item': lb})
            continue
        if f.get('type') in ('select', 'multi') and f.get('opt'):
            for o in f['opt']:
                nm = (o.get('n') or {}).get('en') or o['v']
                boxes.append({'tok': '{{k_%s_%s}}' % (f['k'], o['v']),
                              'label': nm, 'ord': nm})
            continue

        # เกรด 1–5 — กระดาษพิมพ์เป็น ☐ ห้าช่องต่อหนึ่งหัวข้อ แล้วมีช่องเขียนเกรดต่อท้าย
        if f.get('type') == 'grade':
            for n in range(1, (f.get('max') or 5) + 1):
                boxes.append({'tok': '{{k_%s_%s}}' % (f['k'], n),
                              'label': '%s = %d' % ((f.get('label') or {}).get('en') or f['k'], n)})
            continue

        # ตารางให้คะแนน — หนึ่งข้อได้ ☐ เท่าจำนวนคอลัมน์ เรียงซ้ายไปขวาเหมือนกระดาษ
        if f.get('type') == 'checklist':
            opts = f.get('opts') or [{'v': 'S'}, {'v': 'U'}, {'v': 'NA'}]
            for it in f.get('items') or []:
                for o in opts:
                    nm = (o.get('n') or {}).get('en') or o['v']
                    boxes.append({'tok': '{{k_%s_%s_%s}}' % (f['k'], it['id'], o['v']),
                                  'label': '%s · %s' % (it.get('en') or it['id'], nm),
                                  'ord': nm, 'item': it.get('en') or it.get('th') or it['id']})
            continue

        # แถวซ้ำ — ส่งโครงตารางไปให้ฝั่ง Apps Script หาตารางที่หัวคอลัมน์ตรงกัน
        # แล้ววาง token ลงทีละช่อง ดีกว่าโยนเข้ารายการวางมือ 32 บรรทัด
        if f.get('type') == 'table':
            tables.append({
                'k': f['k'],
                'rows': f.get('rows') or 1,
                'label': (f.get('label') or {}).get('en') or f['k'],
                'cols': [{'k': c['k'],
                          'head': (c.get('label') or {}).get('en') or c['k']}
                         for c in f.get('cols') or []]})
            continue

        lab_en, lab_th = (f.get('label') or {}).get('en'), (f.get('label') or {}).get('th')
        en, th = norm(lab_en), norm(lab_th)
        anchor = field_anchor(f)

        # 1) ป้ายตรงเป๊ะ — ลองแบบเก็บวงเล็บก่อน แล้วค่อยแบบตัดวงเล็บ
        hit = next((cellfull[x] for x in (norm_full(lab_en), norm_full(lab_th))
                    if x and x in cellfull and avail.get(cellfull[x], 0) > 0), None)
        if not hit:
            hit = next((cellset[x] for x in (en, th)
                        if x and x in cellset and avail.get(cellset[x], 0) > 0), None)
        if hit and take(hit):
            # เซลล์ที่มีเส้นประอยู่ในตัวเอง ("Aircraft Type: ______") ต้องแทนเส้นประ
            # ในเซลล์นั้น ไม่ใช่เขียนลงเซลล์ข้าง ๆ ซึ่งมักมีป้ายของช่องถัดไปอยู่แล้ว
            # เคยพลาดตรงนี้: ASF หายไปสี่ช่อง (แบบ ทะเบียน วันที่ บทเรียน) แบบเงียบ ๆ
            head = cell_head(hit)
            if head is not None:
                by_cell.append({'cell': hit, 'head': head, 'tok': tok})
            else:
                by_label.append({'label': hit, 'tok': tok})
            continue

        # 2) บรรทัดเส้นประในเซลล์ — "Signature: ______"
        best, bs = None, 0.55
        for i, (lab, _und) in enumerate(lines):
            if i in used_lines: continue
            s = max(score(anchor, lab), score((f.get('label') or {}).get('th') or '', lab))
            if s > bs: best, bs = i, s
        if best is not None:
            lab, und = lines[best]
            used_lines.add(best)
            by_line.append({'label': lab, 'und': len(und), 'tok': tok})
            continue

        # 3) เซลล์ที่ใกล้เคียงพอ — ไม่ใช่ตรงเป๊ะแต่ชัดว่าอันเดียวกัน
        best, bs = None, 0.7
        for orig in cellset.values():
            if avail.get(orig, 0) <= 0: continue
            s = max(score(anchor, orig), score(lab_th or '', orig))
            if s > bs: best, bs = orig, s
        if best and take(best):
            by_label.append({'label': best, 'tok': tok})
            continue

        rec = {'tok': tok, 'label': (f.get('label') or {}).get('en') or f['k'],
               'labelTh': (f.get('label') or {}).get('th') or '',
               'sign': f.get('type') == 'sign'}
        (approval if f.get('_ap') else unmatched).append(rec)

    n_box_docx = sum(c.count('☐') for c in cells)

    ov = OVERRIDES.get(abbr) or {}
    for lab, tok in (ov.get('labels') or {}).items():
        by_label = [b for b in by_label if b['label'] != lab and b['tok'] != tok]
        by_label.append({'label': lab, 'tok': tok})
        # ระบุเองแล้วก็ไม่ต้องค้างอยู่ในรายการวางมืออีก
        unmatched = [u for u in unmatched if u['tok'] != tok]
    # cells: เซลล์ที่มีทั้งป้ายและตัวยึดตำแหน่งอยู่ด้วยกัน ระบุ head เองได้
    # ("Aircraft Reg.:  HS-___" อยากได้ head แค่ "Aircraft Reg.: " เพราะค่ามี HS- อยู่แล้ว)
    for cell, spec in (ov.get('cells') or {}).items():
        tok = spec['tok'] if isinstance(spec, dict) else spec
        head = spec.get('head') if isinstance(spec, dict) else None
        if head is None: head = cell_head(cell) or ''
        by_cell = [b for b in by_cell if b['tok'] != tok and b['cell'] != cell]
        by_cell.append({'cell': cell, 'head': head, 'tok': tok})
        unmatched = [u for u in unmatched if u['tok'] != tok]
    # manual: บังคับให้ต่อท้ายเอกสาร ไม่ต้องพยายามจับคู่กับป้ายบนกระดาษ
    # ใช้เมื่อกระดาษไม่มีที่ให้จริง ๆ และตัวจับคู่ดันไปเจอป้ายที่ "ดูคล้าย" เข้า
    # ของจริง: EFC ช่อง "รายงานผลรายวิชาจาก TrainHub" ถูกจับไปที่ป้าย "Result"
    # ซึ่งเป็นช่องผลสอบ ถ้าช่องนั้นว่างอยู่ ชื่อไฟล์จะไปโผล่เป็นผลสอบเงียบ ๆ
    for tok in (ov.get('manual') or []):
        by_label = [b for b in by_label if b['tok'] != tok]
        by_cell = [b for b in by_cell if b['tok'] != tok]
        by_line = [b for b in by_line if b['tok'] != tok]
        if not any(u['tok'] == tok for u in unmatched):
            src = next((f for f in fields if '{{%s}}' % f['k'] == tok), None)
            lb = (src or {}).get('label') or {}
            unmatched.append({'tok': tok,
                              'label': (lb.get('en') if isinstance(lb, dict) else lb) or tok,
                              'labelTh': (lb.get('th') if isinstance(lb, dict) else '') or ''})

    # skipTables: กระดาษไม่ได้วางตารางนั้นเป็นคอลัมน์แยกแล้ว
    # EFC รวมชื่อวิชากับผลไว้คอลัมน์เดียว ({{s7_N_shown}}) ตามที่เอกสารฉบับแก้ไขทำ
    # ถ้าไม่กัน ตัวสร้างจะพยายามหาหัวคอลัมน์ที่ไม่มีแล้ว แล้วโยน token ทั้งชุด
    # ไปกองในบล็อก "ยังไม่ได้วาง" ท้ายเอกสาร (ของจริงเคยได้ 33 รายการ)
    st = set(ov.get('skipTables') or [])
    if st:
        tables = [t for t in tables if t.get('k') not in st]

    skip = set(ov.get('skip') or [])
    if skip:
        boxes = [b for b in boxes if b['tok'] not in skip]
        unmatched = [u for u in unmatched if u['tok'] not in skip]
    if ov.get('boxes'):
        # เรียงเองตามที่ ☐ อยู่จริงในกระดาษ — ที่ไม่ได้ระบุถือว่าไม่มีช่องให้ติ๊ก
        # เขียนเป็น {"tok": "...", "ord": "ข้อความข้าง ☐ ในกระดาษ"} ได้ด้วย
        # ใช้เมื่อป้ายบนจอสั้นกว่าข้อความในกระดาษมากจนตัวตรวจลำดับจับคู่ไม่ได้
        # การเขียน ord ไว้คือการยืนยันว่า "ช่องนี้คือข้อความนี้" ถ้ากระดาษถูกแก้
        # จนข้อความเปลี่ยน ตัวตรวจจะเตือนทันที ไม่ใช่ปิดเสียงเตือนทิ้ง
        want = ov['boxes']
        known = {b['tok']: b for b in boxes}
        boxes = []
        for w in want:
            tok = w['tok'] if isinstance(w, dict) else w
            rec = dict(known.get(tok) or {'tok': tok, 'label': ''})
            rec['tok'] = tok
            if isinstance(w, dict) and w.get('ord'):
                rec['ord'] = w['ord']
            boxes.append(rec)

    # token ที่เขียนไว้ในเอกสารแล้วด้วยมือ ไม่ต้องให้ตัววางไปหาที่ให้อีก
    # เอกสารบางใบวาง token เองทั้งใบ (STR ฉบับแก้ไข) ถ้าไม่กันออก
    # ตัววางจะไม่เจอที่ว่างให้วาง แล้วโยนทั้งชุดไปกองในบล็อก "ยังไม่ได้วาง" ท้ายเอกสาร
    all_text = '\n'.join(str(c) for c in cells) + '\n' + \
               '\n'.join(x if isinstance(x, str) else ' '.join(map(str, x)) for x in lines)
    already = set(re.findall(r'\{\{[A-Za-z0-9_]+\}\}', all_text))
    if already:
        by_label = [b for b in by_label if b['tok'] not in already]
        by_cell = [b for b in by_cell if b['tok'] not in already]
        by_line = [b for b in by_line if b['tok'] not in already]
        boxes = [b for b in boxes if b['tok'] not in already]
        unmatched = [u for u in unmatched if u['tok'] not in already]

    # ลำดับช่องติ๊กต้องตรงกับกระดาษ ไม่ใช่แค่จำนวนเท่ากัน
    # boxesPartial: กระดาษมี ☐ มากกว่าที่ฟอร์มใช้ และ ☐ ที่เหลือไม่มีใครติ๊กแล้ว
    # (EFC — ตารางรายวิชา 84 ช่องย้ายไปเป็นไฟล์แนบ เหลือ 10 ช่องแรกที่ยังใช้)
    # ยอมให้วางเฉพาะ N ช่องแรกได้ แต่ยังตรวจลำดับเทียบกับกระดาษเหมือนเดิม
    # ถ้าไม่ตรวจ ก็กลับไปเสี่ยงติ๊กเลื่อนช่องแบบที่ตัวกันนี้ตั้งใจกันตั้งแต่แรก
    partial = bool(ov.get('boxesPartial')) and 0 < len(boxes) <= n_box_docx

    order_warn = []
    if len(boxes) == n_box_docx or partial:
        docl = docx_box_labels(dp)
        for i, b in enumerate(boxes):
            if i >= len(docl): break
            # กระดาษวางช่องติ๊กสองแบบ — ☐ ต่อท้ายชื่อตัวเลือก (ข้างข้อความ)
            # หรือ ☐ อยู่ในคอลัมน์ของตัวเองโดยชื่อรายการอยู่หัวแถว (ตารางให้คะแนน)
            # จึงต้องยอมให้ตรงกับอันใดอันหนึ่ง ไม่งั้นตารางให้คะแนนจะเตือนทุกช่อง
            cand = [c for c in (b.get('ord'), b.get('item'), b.get('label')) if c]
            if cand and not any(box_same(c, docl[i]) for c in cand):
                order_warn.append('ช่องที่ %d: เอกสารว่า "%s" · token คือ %s'
                                  % (i + 1, docl[i][:44], b['tok']))

    return {
        'abbr': abbr,
        'docx': 'D-0507-%s-001.docx' % abbr,
        # เลขกำกับที่คาดว่าจะเจอในเอกสาร — ImportTemplate ใช้ตรวจว่าหยิบฉบับถูก
        # findDocx_ ตัดสินด้วยวันแก้ไขล่าสุดอย่างเดียว ถ้ามีชื่อซ้ำใน Drive
        # แล้วใครไปเปิดฉบับเก่าแล้วเซฟ มันจะกลายเป็นตัวใหม่กว่าและถูกหยิบไปทำแม่แบบ
        # โดยไม่มีอะไรเตือน เกิดขึ้นมาแล้วจริงกับ PWR — แม่แบบขึ้น Rev 02 ทั้งที่เอกสารเป็น Rev 03
        'control': (d.get('control') or ''),
        'orderWarn': order_warn,
        'byLabel': by_label,
        'byLine': by_line,
        'byCell': by_cell,
        'boxes': boxes,
        'boxesPartial': partial,
        'tables': tables,
        'boxesInDocx': n_box_docx,
        'approval': approval,
        'manual': unmatched,
    }, unmatched, lines


def main():
    argv = [a for a in sys.argv[1:] if not a.startswith('-')]
    verbose = '-v' in sys.argv
    codes = argv or sorted(
        os.path.basename(p)[:-5]
        for p in glob.glob(os.path.join(HERE, 'formdefs', '*.json'))
        if not os.path.basename(p).startswith('_'))

    maps, report = {}, []
    for a in codes:
        m, note, lines = build(a, verbose)
        if not m:
            report.append('%-6s ข้าม — %s' % (a, note[0])); continue
        maps[a] = m
        auto = len(m['byLabel']) + len(m['byLine']) + len(m['byCell'])
        warn = ''
        if len(m['boxes']) != m['boxesInDocx']:
            warn = '  ⚠️ ช่องติ๊กในนิยามฟอร์ม %d ≠ ☐ ในเอกสาร %d' % (len(m['boxes']), m['boxesInDocx'])
        report.append('%-6s วางอัตโนมัติ %2d (เซลล์ %d · ในเซลล์ %d · เส้นประ %d) · ตาราง %d · ต่อท้ายส่วนอนุมัติ %d · วางมือ %d · ช่องติ๊ก %d%s'
                      % (a, auto, len(m['byLabel']), len(m['byCell']), len(m['byLine']), len(m['tables']),
                         len(m['approval']), len(note), len(m['boxes']), warn))
        if m.get('orderWarn'):
            report.append('       🔴 ลำดับช่องติ๊กไม่ตรงกระดาษ %d จุด — ใบที่ออกมาจะติ๊กผิดข้อ'
                          % len(m['orderWarn']))
            for w in m['orderWarn'][:4]:
                report.append('          ' + w)
        if note:
            report.append('       ต้องวางมือ: ' + ', '.join(x['label'] for x in note[:6])
                          + (' …อีก %d' % (len(note) - 6) if len(note) > 6 else ''))
        if verbose and lines:
            report.append('       เส้นประในเอกสาร: ' + ' · '.join(l for l, _ in lines))

    out = os.path.join(HERE, 'gas', 'TokenMap.gs')

    # ── รวมกับของเดิม ไม่ใช่เขียนทับ ──────────────────────────
    # เดิมสั่งทำใบเดียวแล้วไฟล์เหลือใบเดียว อีก 16 ใบหายเงียบ ๆ
    # ทั้งที่หัวไฟล์นี้เองบอกให้ใช้แบบ  make_tokenmap.py VSR HIF  ได้
    # ตัวส่งออกอ่าน TOKEN_MAP ตัวนี้ ใบที่หายจึงพังด้วยข้อความ
    # "ไม่มี X ใน TokenMap.gs" และรู้ตัวตอนมีคนกดส่งใบจริงเท่านั้น
    if argv and os.path.exists(out):
        try:
            src = open(out, encoding='utf-8').read()
            i = src.index('var TOKEN_MAP = ') + len('var TOKEN_MAP = ')
            keep = json.loads(src[i:src.rindex(';')])
        except Exception as e:
            sys.exit('อ่าน TokenMap.gs เดิมไม่ได้ (%s)\n'
                     'รันแบบไม่ระบุใบเพื่อสร้างใหม่ทั้งไฟล์' % e)
        # ใบที่สั่งทำแล้วสร้างไม่สำเร็จ ต้องหายไปจริง ไม่ใช่ค้างของเก่าไว้
        for a in [x for x in codes if x not in maps]:
            keep.pop(a, None)
        untouched = len([a for a in keep if a not in codes])
        keep.update(maps)
        maps = dict(sorted(keep.items()))
        report.append('รวมกับของเดิม — คงไว้อีก %d ใบที่ไม่ได้สั่งทำรอบนี้' % untouched)

    with open(out, 'w', encoding='utf-8') as fh:
        fh.write('/**\n * TokenMap.gs — สร้างอัตโนมัติจาก tools/make_tokenmap.py\n'
                 ' * อย่าแก้ด้วยมือ · แก้ที่นิยามฟอร์มแล้วรันเครื่องมือใหม่\n'
                 ' *\n'
                 ' * byLabel  ป้ายในเซลล์ -> token ที่ใส่ในเซลล์ถัดไปทางขวา\n'
                 ' * byLine   "ป้าย: ______" ในเซลล์ -> แทนเส้นประด้วย token\n'
                 ' * boxes    token ช่องติ๊ก เรียงตามลำดับที่ ☐ ปรากฏ\n'
                 ' * manual   จับคู่ไม่ได้ ต้องวางมือ — สคริปต์พิมพ์ไว้ท้ายเอกสาร\n */\n')
        fh.write('var TOKEN_MAP = ' + json.dumps(maps, ensure_ascii=False, indent=2) + ';\n')

    print('\n'.join(report))
    print('\nเขียน gas/TokenMap.gs — %d ฟอร์ม' % len(maps))


if __name__ == '__main__':
    main()
