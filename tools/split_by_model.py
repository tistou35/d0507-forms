#!/usr/bin/env python3
# ============================================================
# split_by_model.py — แยกเช็กลิสต์ฉบับรวมออกเป็นฉบับต่อรุ่นอากาศยาน
#
#   python3 tools/split_by_model.py NPC EPC
#
# ── ทำไม ─────────────────────────────────────────────────────
# D-0507-NPC-001 ฉบับเดียวครอบ C172 M/N/S มาตลอด สามรุ่นจึงแก้แยกกันไม่ได้
# แยกเป็นสามฉบับ ตอนนี้เนื้อหาเหมือนกันทุกตัวอักษร แต่ต่อไปแก้ทีละรุ่นได้
#
# ── เลขเอกสาร ────────────────────────────────────────────────
# แยกด้วยตัวย่อรุ่น ตามแบบ PCR-FI / PCR-TKI ที่ใช้อยู่แล้วในทะเบียน
#   D-0507-NPC-M-001  NP-CHK-303-A   C172M
#   D-0507-NPC-N-001  NP-CHK-304-A   C172N
#   D-0507-NPC-S-001  NP-CHK-305-A   C172S
# เป็นเอกสารใหม่ทั้งสามฉบับ จึงเริ่มที่ตัวอักษร -A และฉบับ 01 / แก้ไข 00
# ฉบับรวมเดิมถูกแทนที่ ไม่ได้ถูกแก้ — เก็บไว้เป็น *-series-archive.docx
# ============================================================
import os, shutil, sys

import docx

HERE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SRC_DIR = os.path.dirname(HERE)

MODELS = [('M', 'C172M', 303), ('N', 'C172N', 304), ('S', 'C172S', 305)]
PREFIX = {'NPC': 'NP-CHK', 'EPC': 'EP-CHK'}


def retext(doc, pairs):
    """แทนคำทีละ run — ไม่รวม run เข้าด้วยกัน รูปแบบตัวอักษรจึงไม่เพี้ยน

    แทนที่ทั้งย่อหน้ารวดเดียวจะทำให้ตัวหนา/ขนาดที่ต่างกันในบรรทัดเดียวหายไป
    หัวกระดาษของเอกสารพวกนี้มีทั้งตัวหนาและตัวธรรมดาอยู่ในย่อหน้าเดียวกัน
    """
    n = 0
    for p in doc.paragraphs:
        for r in p.runs:
            for a, b in pairs:
                if a in r.text:
                    r.text = r.text.replace(a, b)
                    n += 1
    return n


def build(abbr):
    src = os.path.join(SRC_DIR, 'D-0507-%s-001.docx' % abbr)
    if not os.path.exists(src):
        sys.exit('ไม่พบ %s' % src)
    old_code = None
    d0 = docx.Document(src)
    for p in d0.paragraphs:
        if PREFIX[abbr] + '-303-' in p.text:
            i = p.text.index(PREFIX[abbr])
            old_code = p.text[i:i + len(PREFIX[abbr]) + 6]
            break
    if not old_code:
        sys.exit('%s: หาเลขกำกับเดิมในเอกสารไม่เจอ' % abbr)

    made = []
    for letter, model, seq in MODELS:
        out = os.path.join(SRC_DIR, 'D-0507-%s-%s-001.docx' % (abbr, letter))
        shutil.copy2(src, out)
        d = docx.Document(out)
        new_code = '%s-%d-A' % (PREFIX[abbr], seq)
        hits = retext(d, [
            (old_code, new_code),
            ('Cessna 172 Series', 'Cessna ' + model[1:]),   # C172M → 172M
            ('Cessna 172 ', 'Cessna %s ' % model[1:]),      # ท้ายกระดาษ
        ])
        d.save(out)
        made.append((model, new_code, os.path.basename(out), hits))

    arch = os.path.join(SRC_DIR, 'D-0507-%s-001-series-archive.docx' % abbr)
    if not os.path.exists(arch):
        shutil.move(src, arch)
    print('✅ %s → %d ฉบับ (ฉบับรวมเดิมเก็บเป็น %s)' % (abbr, len(made), os.path.basename(arch)))
    for m, c, f, h in made:
        print('   %-6s %-14s %-28s แทนคำ %d จุด' % (m, c, f, h))


if __name__ == '__main__':
    if len(sys.argv) < 2:
        sys.exit('ใช้: python3 tools/split_by_model.py NPC [EPC]')
    for a in sys.argv[1:]:
        build(a)
