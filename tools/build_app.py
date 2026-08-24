#!/usr/bin/env python3
# ============================================================
# build_app.py — เขียน D-0507-APP-001.docx ใหม่ให้ตรงกับฟอร์มที่ใช้จริง
#
#   python3 tools/build_app.py
#
# ── ทำไม ─────────────────────────────────────────────────────
# ฉบับกระดาษเดิม (AD-APP-302-B) เป็นใบสมัครเข้าเรียน/ขอ rating — ถามหลักสูตร
# ที่จะเรียน ชั่วโมงบิน และใบอนุญาตนักบิน
# ส่วนฟอร์มที่ใช้จริงที่ AeroFBO › recruitment.php เป็นใบสมัครงาน — ถามตำแหน่ง
# ที่สมัคร ประวัติการศึกษา และใบรับรองที่ถืออยู่
# คนกรอกกระดาษกับคนกรอกเว็บจึงตอบคนละชุดคำถาม เทียบกันไม่ได้
#
# ── ยึดอะไรเป็นหลัก ──────────────────────────────────────────
# โครงและชื่อช่องยึดตามหน้าเว็บจริง (อ่านจาก DOM ของ recruitment.php)
# รูปลักษณ์ยึดตามฉบับเดิม — แถบสีกรม ป้ายเทาอ่อน ช่องกรอกขาว บล็อกคำรับรองฟ้าอ่อน
# เพื่อให้ยังเป็นเอกสารชุดเดียวกับใบอื่นของ D-0507
# ============================================================
import os

import docx
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

HERE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT = os.path.join(os.path.dirname(HERE), 'D-0507-APP-001.docx')
ARCH = os.path.join(os.path.dirname(HERE), 'D-0507-APP-001-enrolment-archive.docx')

CODE, EFF = 'AD-APP-302-C', '24 AUG 2026'
NAVY, MID = '1F3864', '2E75B6'
LBL_BG, DECL_BG, HEAD_BG = 'F5F5F5', 'D6E4F7', 'F5F5F5'
GREY, INK = '777777', '1A1A1A'

# ตำแหน่งที่รับสมัคร — ตรงกับ <select name="role_apply"> ในหน้าเว็บ
ROLES = ['Administrator', 'Maintenance Manager', 'QA Manager', 'Flight Attendant',
         'Pilot', 'Student Pilot', 'Mechanic', 'General Staff',
         'AE Instructor', 'Instructor']
EDU_LEVELS = "High School  ·  Bachelor's Degree  ·  Master's Degree  ·  Doctorate"
CERT_TYPES = 'Licence  ·  Certificate  ·  Endorsement  ·  Medical  ·  Identification'
DECLARE = ('I hereby certify that the information provided in this application is true and '
           'complete. I understand that any false statements may result in disqualification.')


def shade(cell, hexfill):
    el = OxmlElement('w:shd')
    el.set(qn('w:val'), 'clear'); el.set(qn('w:fill'), hexfill)
    cell._tc.get_or_add_tcPr().append(el)


def write(cell, text, size=8.5, bold=False, color=INK, italic=False, align=None):
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.font.name = 'Arial'; r.font.size = Pt(size); r.bold = bold; r.italic = italic
    r.font.color.rgb = RGBColor.from_string(color)
    return p


def borders(tbl, color='BFBFBF'):
    pr = tbl._tbl.tblPr
    b = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement('w:' + side)
        e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), '4'); e.set(qn('w:color'), color)
        b.append(e)
    pr.append(b)


def no_borders(tbl):
    pr = tbl._tbl.tblPr
    b = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement('w:' + side); e.set(qn('w:val'), 'nil'); b.append(e)
    pr.append(b)


WIDTH = 178      # A4 210mm ลบขอบซ้ายขวาข้างละ 16mm


def add_table(doc, rows, cols, widths=None):
    t = doc.add_table(rows=rows, cols=cols)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    # ไม่ระบุความกว้างมา = แบ่งเท่ากันเต็มหน้า · ปล่อยให้ autofit จัดเอง
    # แถบหัวข้อจะหดตามความยาวข้อความจนสั้นกุด
    widths = widths or [WIDTH / cols] * cols
    # ต้องตั้งความกว้างที่ตัวตาราง ไม่ใช่แค่ที่ช่อง — ตั้งแต่ช่องอย่างเดียว
    # Word ยังคำนวณใหม่ตามความยาวข้อความ แถบหัวข้อจึงหดสั้นกุด
    pr = t._tbl.tblPr
    w = OxmlElement('w:tblW')
    w.set(qn('w:w'), str(int(WIDTH * 56.7))); w.set(qn('w:type'), 'dxa')
    pr.append(w)
    lay = OxmlElement('w:tblLayout'); lay.set(qn('w:type'), 'fixed'); pr.append(lay)
    for r in t.rows:
        for c, wd in zip(r.cells, widths):
            c.width = Mm(wd)
    borders(t)
    return t


def gap(doc, pt=3):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(pt)
    return p


def field_block(doc, rows, total=178):
    """หนึ่งบล็อก = แถวป้ายเทา สลับกับแถวช่องกรอกขาว

    rows = [[(ป้าย, สัดส่วนกว้าง), …], …] หนึ่งรายการต่อหนึ่งคู่แถว
    """
    ncol = max(len(r) for r in rows)
    t = add_table(doc, len(rows) * 2, ncol)
    for i, row in enumerate(rows):
        span = ncol // len(row)
        for j, (lab, _) in enumerate(row):
            a = j * span
            b = (a + span - 1) if j < len(row) - 1 else ncol - 1
            lc = t.rows[i * 2].cells[a]
            ic = t.rows[i * 2 + 1].cells[a]
            if b > a:
                lc = lc.merge(t.rows[i * 2].cells[b])
                ic = ic.merge(t.rows[i * 2 + 1].cells[b])
            shade(lc, LBL_BG); write(lc, lab, 7.5, True, GREY)
            shade(ic, 'FFFFFF'); write(ic, ' ')
            ic.paragraphs[0].paragraph_format.space_after = Pt(6)
    return t


def grid_block(doc, headers, nrows):
    """ตารางรายการที่กรอกได้หลายแถว (ประวัติการศึกษา · ใบรับรอง)"""
    t = add_table(doc, nrows + 1, len(headers))
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]
        shade(c, LBL_BG); write(c, h, 7.5, True, GREY)
    for i in range(1, nrows + 1):
        row_bg = 'FFFFFF' if i % 2 else 'FAFAFA'
        for j in range(len(headers)):
            shade(t.rows[i].cells[j], row_bg)
            write(t.rows[i].cells[j], ' ')
            t.rows[i].cells[j].paragraphs[0].paragraph_format.space_after = Pt(6)
    return t


def band(doc, text, sub=None):
    t = add_table(doc, 2 if sub else 1, 1)
    no_borders(t)
    c = t.rows[0].cells[0]
    shade(c, NAVY); write(c, text, 11, True, 'FFFFFF')
    if sub:
        c2 = t.rows[1].cells[0]
        shade(c2, MID); write(c2, sub, 8, False, 'FFFFFF')
    hairline(doc)
    return t


def hairline(doc):
    """ย่อหน้าบางที่สุดเท่าที่ทำได้ ไว้คั่นตารางสองตัวไม่ให้ถูกรวมเป็นตัวเดียว

    Word รวมตารางที่ติดกันโดยไม่มีย่อหน้าคั่น แถบหัวข้อจึงไปได้ความกว้าง
    เท่าคอลัมน์แรกของตารางถัดไป กลายเป็นแถบสั้นกุดและตัวหนังสือตัดบรรทัด
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(1)
    r = p.add_run('')
    r.font.size = Pt(1)
    return p


def section(doc, text):
    t = add_table(doc, 1, 1)
    no_borders(t)
    c = t.rows[0].cells[0]
    shade(c, NAVY); write(c, text, 8.5, True, 'FFFFFF')
    hairline(doc)
    return t


def build():
    doc = docx.Document()
    s = doc.sections[0]
    s.page_width, s.page_height = Mm(210), Mm(297)
    s.left_margin = s.right_margin = s.top_margin = s.bottom_margin = Mm(16)
    st = doc.styles['Normal']
    st.font.name = 'Arial'; st.font.size = Pt(8.5)

    # ── หัวเอกสาร ──
    t = add_table(doc, 1, 2, [104, 74])
    no_borders(t)
    c = t.rows[0].cells[0]
    shade(c, HEAD_BG)
    p = write(c, 'D-0507 Flight Training Co., Ltd.', 10, True, NAVY)
    r = p.add_run('\nApproved Training Organisation (ATO)')
    r.font.name = 'Arial'; r.font.size = Pt(7.5); r.font.color.rgb = RGBColor.from_string(GREY)
    c2 = t.rows[0].cells[1]
    p = write(c2, CODE + '    EFF: ' + EFF, 7.5, False, GREY,
              align=WD_ALIGN_PARAGRAPH.RIGHT)
    r = p.add_run('\nAdministration (AD)  ·  D-0507 ATO')
    r.font.name = 'Arial'; r.font.size = Pt(7); r.italic = True
    r.font.color.rgb = RGBColor.from_string(GREY)
    gap(doc)

    band(doc, 'APPLICATION FORM',
         'Personnel application — complete all sections  ·  '
         'Submitted online at AeroFBO › Recruitment')
    gap(doc)

    # ── ตำแหน่งที่สมัคร ──
    section(doc, 'ROLE APPLIED FOR')
    t = add_table(doc, 2, 5)
    for i, role in enumerate(ROLES):
        c = t.rows[i // 5].cells[i % 5]
        shade(c, 'FFFFFF'); write(c, '☐  ' + role, 7.5)
    gap(doc)

    # ── ข้อมูลส่วนตัว ──
    section(doc, 'PERSONAL INFORMATION')
    field_block(doc, [
        [('First Name:', 1), ('Last Name:', 1)],
        [('Email Address:', 1), ('Phone Number:', 1)],
        [('Date of Birth:', 1), ('Nationality:', 1)],
        [('ID Card / Passport No.:', 1), ('Profile Photo attached:  ☐ Yes   ☐ No', 1)],
        [('Permanent Address:', 1)],
    ])
    gap(doc)

    # ── ประวัติการศึกษา ──
    section(doc, 'EDUCATION HISTORY')
    grid_block(doc, ['Education Level', 'Institution Name', 'Major / Field',
                     'Year Start', 'Year Finish'], 4)
    p = doc.add_paragraph()
    r = p.add_run('Education level: ' + EDU_LEVELS)
    r.font.name = 'Arial'; r.font.size = Pt(6.5); r.italic = True
    r.font.color.rgb = RGBColor.from_string(GREY)
    p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(4)

    # ── ใบอนุญาตและใบรับรอง ──
    section(doc, 'LICENCES & CERTIFICATIONS')
    grid_block(doc, ['Document Type', 'Licence / Certificate Name', 'Certificate No.',
                     'Issuing Authority', 'Issue Date', 'Expiry Date'], 4)
    p = doc.add_paragraph()
    r = p.add_run('Document type: ' + CERT_TYPES
                  + '     ·     Attach a copy of each document listed above.')
    r.font.name = 'Arial'; r.font.size = Pt(6.5); r.italic = True
    r.font.color.rgb = RGBColor.from_string(GREY)
    p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(4)

    # ── ประวัติการทำงาน / ความสามารถ ──
    section(doc, 'PROFESSIONAL PROFILE')
    t = add_table(doc, 1, 1)
    c = t.rows[0].cells[0]
    shade(c, 'FFFFFF'); write(c, ' ')
    c.paragraphs[0].paragraph_format.space_after = Pt(20)
    gap(doc)

    # ── คำรับรองและลายเซ็น ──
    section(doc, 'DECLARATION & SIGNATURE')
    t = add_table(doc, 1, 1)
    c = t.rows[0].cells[0]
    shade(c, DECL_BG); write(c, DECLARE, 7.5, False, GREY, italic=True)
    gap(doc, 2)

    t = add_table(doc, 2, 3, [70, 54, 54])
    for j, lab in enumerate(['Applicant Signature:', 'Date:', 'Received by:']):
        c = t.rows[0].cells[j]
        shade(c, LBL_BG); write(c, lab, 7.5, True, NAVY)
        c2 = t.rows[1].cells[j]
        shade(c2, 'FFFFFF'); write(c2, ' ')
        c2.paragraphs[0].paragraph_format.space_after = Pt(7)

    # ── ท้ายกระดาษ ──
    ftr = doc.sections[0].footer
    fp = ftr.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = fp.add_run('D-0507 Flight Training Co., Ltd.  |  APPLICATION FORM  |  '
                   + CODE + '  EFF: ' + EFF + '  |  Controlled Copy — Uncontrolled when printed')
    r.font.name = 'Arial'; r.font.size = Pt(6.5)
    r.font.color.rgb = RGBColor.from_string(GREY)

    if os.path.exists(OUT) and not os.path.exists(ARCH):
        os.rename(OUT, ARCH)
        print('เก็บฉบับใบสมัครเข้าเรียนเดิมไว้ที่', os.path.basename(ARCH))
    doc.save(OUT)
    print('เขียน', os.path.basename(OUT), '·', CODE, '· EFF', EFF)


if __name__ == '__main__':
    build()
