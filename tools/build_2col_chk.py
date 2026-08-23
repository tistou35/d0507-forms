#!/usr/bin/env python3
# ============================================================
# build_2col_chk.py — จัดเช็กลิสต์ห้องนักบิน (NPC / EPC) ให้เป็นสองคอลัมน์
#
#   python3 tools/build_2col_chk.py NPC          สองคอลัมน์ (ค่าปริยาย)
#   python3 tools/build_2col_chk.py NPC EPC
#   python3 tools/build_2col_chk.py --cols 3 NPC EPC
#
# ── ทำไม ─────────────────────────────────────────────────────
# ตารางเดิมเป็นคอลัมน์เดียวยาว 172 แถว (NPC) / 112 แถว (EPC) ออกมา 4 และ 3 หน้า
# เช็กลิสต์ที่ใช้ในห้องนักบินต้องพลิกให้น้อยที่สุด สองคอลัมน์ลดลงเหลือครึ่งหนึ่ง
#
# ── วิธี ─────────────────────────────────────────────────────
# ไม่ได้พิมพ์เนื้อหาใหม่ — คัดลอก <w:tc> เดิมทั้งก้อนไปวางในตารางใหม่
# ฟอนต์ เงาพื้น ช่อง ☐ และการจัดย่อหน้าจึงเหมือนเดิมทุกเซลล์
# เขียนเนื้อหาใหม่เองเมื่อไรคือเปิดโอกาสให้พิมพ์ผิดในเอกสารที่ใช้บนเครื่องบิน
#
# แบ่งครึ่งตามกลุ่มหัวข้อ ไม่ใช่ตามจำนวนแถว — หัวข้อกับรายการใต้มันต้องอยู่
# คอลัมน์เดียวกัน ไม่งั้นอ่านเจอ "☐ Mixture | IDLE CUT-OFF" ลอยอยู่หัวคอลัมน์
# โดยไม่รู้ว่าเป็นขั้นตอนของเหตุฉุกเฉินอะไร
#
# ── เอกสารควบคุม ─────────────────────────────────────────────
# รูปแบบที่พิมพ์ออกมาเปลี่ยน = ฉบับพิมพ์เปลี่ยน จึงต้องเลื่อนตัวอักษรท้ายเลขกำกับ
# (-B → -C) ตามกฎใน CLAUDE.md และเก็บฉบับก่อนแก้ไว้ย้อนดู
# ============================================================
import copy, os, shutil, sys

import docx
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SRC_DIR = os.path.dirname(HERE)                      # โฟลเดอร์ Manual revision

# ความกว้างรวมต้องเท่าของเดิม 10106 dxa (1 dxa = 1/20 pt)
TOTAL, W_GAP = 10106, 170
# สัดส่วนช่องรายการ : ช่องสิ่งที่ต้องทำ
# ฉบับเดิมคอลัมน์เดียวใช้ 0.62 แต่พอแบ่งสามคอลัมน์ ช่องขวาจะแคบจนคำอย่าง
# UNOBSTRUCTED ถูกตัดกลางคำเป็น "UNOBSTRUCT ED" — อ่านผิดได้ในห้องนักบิน
# ยิ่งคอลัมน์เยอะยิ่งต้องแบ่งให้ช่องขวามากขึ้น
RATIO = {1: 0.62, 2: 0.615, 3: 0.55}


def rows_of(tbl):
    """แยกแถวเป็นกลุ่ม: หัวข้อหนึ่งอันพร้อมรายการใต้มัน"""
    groups, cur = [], None
    for tr in tbl.rows:
        tcs = tr._tr.findall(qn('w:tc'))
        # แถวหัวข้อคือแถวที่ผสานสองช่องเป็นช่องเดียว
        is_sec = len(tcs) == 1 or tr.cells[0]._tc is tr.cells[1]._tc
        if is_sec or cur is None:
            cur = {'sec': tr if is_sec else None, 'items': []}
            groups.append(cur)
            if is_sec:
                continue
        cur['items'].append(tr)
    return groups


def split(groups, cols=2):
    """ตัดให้ทุกคอลัมน์ยาวใกล้กันที่สุด โดยไม่ฉีกกลุ่มหัวข้อ

    ไล่ตัดทีละจุดตามลำดับ ไม่ได้จัดเรียงใหม่ — ลำดับขั้นตอนของเช็กลิสต์
    ต้องอ่านจากบนลงล่างคอลัมน์ซ้ายให้จบก่อนแล้วขึ้นหัวคอลัมน์ถัดไป
    """
    size = [1 + len(g['items']) if g['sec'] is not None else len(g['items']) for g in groups]
    total = sum(size)
    out, start, done = [], 0, 0
    for c in range(cols - 1):
        target = total * (c + 1) / cols
        n, cut, best = done, start, float('inf')
        for i in range(start, len(groups) + 1):
            if abs(n - target) < best:
                best, cut = abs(n - target), i
            if i < len(groups):
                n += size[i]
        cut = max(cut, start)
        out.append(groups[start:cut])
        done += sum(size[start:cut])
        start = cut
    out.append(groups[start:])
    return out


def flat(groups):
    """คลี่กลุ่มกลับเป็นลำดับแถว — (tr, เป็นหัวข้อไหม)"""
    out = []
    for g in groups:
        if g['sec'] is not None:
            out.append((g['sec'], True))
        out += [(tr, False) for tr in g['items']]
    return out


def blank_tc(w):
    """ช่องคั่นกลาง ไม่มีเส้นขอบ"""
    tc = OxmlElement('w:tc')
    pr = OxmlElement('w:tcPr')
    tcw = OxmlElement('w:tcW'); tcw.set(qn('w:w'), str(w)); tcw.set(qn('w:type'), 'dxa')
    pr.append(tcw)
    bd = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        e = OxmlElement('w:' + side); e.set(qn('w:val'), 'nil'); bd.append(e)
    pr.append(bd)
    tc.append(pr)
    tc.append(OxmlElement('w:p'))
    return tc


def cells_from(tr, is_sec, width_item, width_act):
    """คัดลอกช่องจากแถวเดิม คืนมาเป็นสองช่องเสมอ (หัวข้อ = ผสานสองช่อง)"""
    tcs = tr._tr.findall(qn('w:tc'))
    if is_sec:
        tc = copy.deepcopy(tcs[0])
        set_w(tc, width_item + width_act)
        span(tc, 2)
        return [tc]
    a = copy.deepcopy(tcs[0]); set_w(a, width_item); span(a, None)
    b = copy.deepcopy(tcs[1] if len(tcs) > 1 else tcs[0]); set_w(b, width_act); span(b, None)
    return [a, b]


def set_w(tc, w):
    pr = tc.find(qn('w:tcPr'))
    if pr is None:
        pr = OxmlElement('w:tcPr'); tc.insert(0, pr)
    old = pr.find(qn('w:tcW'))
    if old is not None:
        pr.remove(old)
    e = OxmlElement('w:tcW'); e.set(qn('w:w'), str(w)); e.set(qn('w:type'), 'dxa')
    pr.insert(0, e)


def span(tc, n):
    pr = tc.find(qn('w:tcPr'))
    old = pr.find(qn('w:gridSpan'))
    if old is not None:
        pr.remove(old)
    if n:
        e = OxmlElement('w:gridSpan'); e.set(qn('w:val'), str(n)); pr.append(e)


def empty_pair(width_item, width_act):
    return [blank_tc(width_item), blank_tc(width_act)]


def widths(cols):
    half = (TOTAL - W_GAP * (cols - 1)) // cols
    item = int(half * RATIO.get(cols, 0.55))
    return item, half - item


def build(abbr, cols=2):
    W_ITEM, W_ACT = widths(cols)
    src = os.path.join(SRC_DIR, 'D-0507-%s-001.docx' % abbr)
    arch = os.path.join(SRC_DIR, 'D-0507-%s-001-1col-archive.docx' % abbr)
    d = docx.Document(src)
    tbl = d.tables[-1]

    groups = rows_of(tbl)
    parts = [flat(g) for g in split(groups, cols)]
    if any(not p for p in parts):
        sys.exit('%s: แบ่ง %d คอลัมน์ไม่ได้ — กลุ่มหัวข้อน้อยเกินไป' % (abbr, cols))

    # ── ตารางใหม่: item | action | คั่น | item | action ──
    new = copy.deepcopy(tbl._tbl)
    for tr in new.findall(qn('w:tr')):
        new.remove(tr)
    grid = new.find(qn('w:tblGrid'))
    for gc in grid.findall(qn('w:gridCol')):
        grid.remove(gc)
    for n in range(cols):
        if n:
            gc = OxmlElement('w:gridCol'); gc.set(qn('w:w'), str(W_GAP)); grid.append(gc)
        for w in (W_ITEM, W_ACT):
            gc = OxmlElement('w:gridCol'); gc.set(qn('w:w'), str(w)); grid.append(gc)

    n_rows = max(len(p) for p in parts)
    for i in range(n_rows):
        tr = OxmlElement('w:tr')
        for n, part in enumerate(parts):
            if n:
                tr.append(blank_tc(W_GAP))
            if i < len(part):
                for tc in cells_from(part[i][0], part[i][1], W_ITEM, W_ACT):
                    tr.append(tc)
            else:
                for tc in empty_pair(W_ITEM, W_ACT):
                    tr.append(tc)
        new.append(tr)

    tbl._tbl.addnext(new)
    tbl._tbl.getparent().remove(tbl._tbl)

    # ── เลื่อนตัวอักษรท้ายเลขกำกับ — ฉบับพิมพ์เปลี่ยน ──
    bumped = set()
    for p in d.paragraphs:
        for r in p.runs:
            if '-CHK-303-B' in r.text:
                r.text = r.text.replace('-CHK-303-B', '-CHK-303-C')
                bumped.add(p.text[:40])

    if os.path.exists(arch):
        sys.exit('%s: มีไฟล์เก็บฉบับเดิมอยู่แล้ว — ลบหรือเปลี่ยนชื่อก่อน\n  %s' % (abbr, arch))
    shutil.copy2(src, arch)
    d.save(src)
    print('✅ %-4s %d แถว → %d แถว × %d คอลัมน์ (%s) · เลื่อนเลขกำกับ %d จุด'
          % (abbr, len(tbl.rows), n_rows, cols,
             ' · '.join(str(len(p)) for p in parts), len(bumped)))
    print('   เก็บฉบับคอลัมน์เดียวไว้ที่ %s' % os.path.basename(arch))


if __name__ == '__main__':
    args = sys.argv[1:]
    cols = 2
    if args and args[0] == '--cols':
        cols = int(args[1]); args = args[2:]
    if not args:
        sys.exit('ใช้: python3 tools/build_2col_chk.py [--cols N] NPC [EPC]')
    for a in args:
        build(a, cols)
