#!/usr/bin/env python3
# ============================================================
# build_c206.py — สร้างเช็กลิสต์ Cessna T206H สองใบ
#
#   python3 tools/build_c206.py
#
# ── วิธี ─────────────────────────────────────────────────────
# ไม่ได้ทำแบบฟอร์มขึ้นใหม่ — คัดลอกโครงจากใบของ C172S ที่ใช้อยู่ แล้วเปลี่ยน
# เฉพาะแถวในตาราง หัวกระดาษ เงาพื้นของแถวหัวข้อ ฟอนต์ และช่อง ☐ จึงเหมือนกัน
# ทุกจุดกับใบที่ครูการบินถืออยู่แล้ว ไม่ต้องเรียนรู้หน้าตาใหม่
#
# เนื้อหาอยู่ใน c206_content.py ถอดจาก Cessna Pilot's Checklist T206H
# (Rev 4 · 13 May 2002) ที่ผู้ใช้ส่งมาเป็นภาพถ่าย
#
# ⚠️ ภาพถ่ายต้นฉบับถูกตัดท้ายหน้าเกือบทุกหน้า รายการที่หายทำเป็นแถบเตือน
#    สีแดงบนกระดาษ ไม่ได้เดาเนื้อหาแทน — จนกว่าจะเติมครบ ใบนี้เป็นฉบับร่าง
# ============================================================
import copy
import os
import sys

import docx
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import RGBColor

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from build_2col_chk import (blank_tc, cells_from, empty_pair, flat, rows_of,  # noqa: E402
                            split, widths, W_GAP)
from c206_content import EMERGENCY, GAP, NORMAL  # noqa: E402

HERE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SRC_DIR = os.path.dirname(HERE)
EFF = '25 AUG 2026'
MODEL = 'T206H'

JOBS = [
    # (ตัวย่อ, ต้นแบบ, เนื้อหา, เลขกำกับใหม่, จำนวนคอลัมน์)
    ('NPC-T206H', 'D-0507-NPC-S-001.docx', NORMAL,    'NP-CHK-306-A', 2),
    ('EPC-T206H', 'D-0507-EPC-S-001.docx', EMERGENCY, 'EP-CHK-306-A', 2),
]


def set_cell_text(tc, text, bold=False, red=False):
    """เขียนข้อความลงเซลล์โดยคง run แรกไว้ — รูปแบบตัวอักษรของต้นแบบจึงติดมาด้วย"""
    ps = tc.findall(qn('w:p'))
    keep = ps[0]
    for p in ps[1:]:
        tc.remove(p)
    runs = keep.findall(qn('w:r'))
    if not runs:
        r = OxmlElement('w:r'); keep.append(r); runs = [r]
    for r in runs[1:]:
        keep.remove(r)
    r = runs[0]
    for t in r.findall(qn('w:t')):
        r.remove(t)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    if bold or red:
        pr = r.find(qn('w:rPr'))
        if pr is None:
            pr = OxmlElement('w:rPr'); r.insert(0, pr)
        if bold and pr.find(qn('w:b')) is None:
            pr.append(OxmlElement('w:b'))
        if red:
            for old in pr.findall(qn('w:color')):
                pr.remove(old)
            c = OxmlElement('w:color'); c.set(qn('w:val'), 'C00000'); pr.append(c)


def one_col_table(tbl, data):
    """แทนแถวทั้งหมดด้วยเนื้อหาใหม่ ใช้แถวเดิมเป็นแม่พิมพ์

    ต้องมีทั้งแม่พิมพ์ของแถวหัวข้อ (ผสานสองช่อง) และของแถวรายการ
    ถ้าใช้แถวรายการไปทำหัวข้อ เงาพื้นสีเข้มจะหายไปทั้งใบ
    """
    proto_sec, proto_item = None, None
    for tr in tbl.rows:
        tcs = tr._tr.findall(qn('w:tc'))
        is_sec = len(tcs) == 1 or tr.cells[0]._tc is tr.cells[1]._tc
        if is_sec and proto_sec is None:
            proto_sec = copy.deepcopy(tr._tr)
        if not is_sec and proto_item is None:
            proto_item = copy.deepcopy(tr._tr)
        if proto_sec is not None and proto_item is not None:
            break
    if proto_sec is None or proto_item is None:
        sys.exit('ต้นแบบไม่มีทั้งแถวหัวข้อและแถวรายการ')

    body = tbl._tbl
    for tr in body.findall(qn('w:tr')):
        body.remove(tr)

    n_sec = n_item = n_gap = 0
    for title, items in data:
        tr = copy.deepcopy(proto_sec)
        set_cell_text(tr.findall(qn('w:tc'))[0], title)
        body.append(tr)
        n_sec += 1
        for item, act in items:
            tr = copy.deepcopy(proto_item)
            tcs = tr.findall(qn('w:tc'))
            if item == GAP:
                set_cell_text(tcs[0], '⚠  ต้นฉบับขาด — ' + act, red=True)
                set_cell_text(tcs[1], 'เติมจาก POH', red=True, bold=True)
                n_gap += 1
            else:
                set_cell_text(tcs[0], '☐  ' + item)
                set_cell_text(tcs[1], act, bold=True)
                n_item += 1
            body.append(tr)
    return n_sec, n_item, n_gap


def relayout(d, cols):
    W_ITEM, W_ACT = widths(cols)
    tbl = d.tables[-1]
    parts = [flat(g) for g in split(rows_of(tbl), cols)]
    if any(not p for p in parts):
        sys.exit('แบ่ง %d คอลัมน์ไม่ได้' % cols)
    new = copy.deepcopy(tbl._tbl)
    for tr in new.findall(qn('w:tr')):
        new.remove(tr)
    grid = new.find(qn('w:tblGrid'))
    for gc in grid.findall(qn('w:gridCol')):
        grid.remove(gc)
    for n in range(cols):
        if n:
            gc = OxmlElement('w:gridCol'); gc.set(qn('w:w'), str(W_GAP)); grid.append(gc)
        for w in (W_ITEM, W_ACT):
            gc = OxmlElement('w:gridCol'); gc.set(qn('w:w'), str(w)); grid.append(gc)
    for i in range(max(len(p) for p in parts)):
        tr = OxmlElement('w:tr')
        for n, part in enumerate(parts):
            if n:
                tr.append(blank_tc(W_GAP))
            src = cells_from(part[i][0], part[i][1], W_ITEM, W_ACT) if i < len(part) \
                else empty_pair(W_ITEM, W_ACT)
            for tc in src:
                tr.append(tc)
        new.append(tr)
    tbl._tbl.addnext(new)
    tbl._tbl.getparent().remove(tbl._tbl)
    return len(new.findall(qn('w:tr')))


def stamp(d, code, gaps):
    """หัวและท้ายกระดาษ + คำเตือนว่ายังเป็นฉบับร่าง"""
    import re
    pat = re.compile(r'(NP|EP)-CHK-\d{3}-[A-Z]')
    hits = 0
    for p in d.paragraphs:
        for r in p.runs:
            t = pat.sub(code, r.text)
            t = re.sub(r'EFF: \d{2} [A-Z]{3} \d{4}', 'EFF: ' + EFF, t)
            t = t.replace('Cessna 172S', 'Cessna ' + MODEL)
            t = t.replace('Cessna 172', 'Cessna ' + MODEL)
            if t != r.text:
                r.text = t
                hits += 1
    # คำเตือนแทนที่ย่อหน้าเดิม — วางไว้ที่เดียวกับคำเตือนของใบ C172
    for p in d.paragraphs:
        if p.text.strip().startswith(('⚠  WARNING', '⚠  CAUTION')):
            for r in p.runs[1:]:
                r.text = ''
            p.runs[0].text = (
                '⚠  ฉบับร่าง — ยังไม่อนุมัติให้ใช้ในการบิน · ถอดจาก Cessna Pilot\'s '
                'Checklist Model T206H (Rev 4 · 13 May 2002) ซึ่งภาพถ่ายต้นฉบับถูกตัด'
                'ท้ายหน้า ทำให้ยังขาดรายการ %d จุดที่ทำแถบสีแดงไว้ '
                'ต้องเติมจาก POH และให้หัวหน้าครูการบินรับรองก่อนใช้งาน' % gaps)
            for r in p.runs:
                r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
            hits += 1
            break
    return hits


def build(abbr, proto, data, code, cols):
    src = os.path.join(SRC_DIR, proto)
    if not os.path.exists(src):
        sys.exit('ไม่พบต้นแบบ %s' % src)
    out = os.path.join(SRC_DIR, 'D-0507-%s-001.docx' % abbr)

    d = docx.Document(src)
    n_sec, n_item, n_gap = one_col_table(d.tables[-1], data)
    rows = relayout(d, cols)
    hits = stamp(d, code, n_gap)
    d.save(out)
    print('  %-10s %-14s %d หมวด · %d รายการ · จุดที่ต้นฉบับขาด %d · '
          '%d คอลัมน์ × %d แถว · หัว/ท้าย %d'
          % (abbr, code, n_sec, n_item, n_gap, cols, rows, hits))
    return n_gap


if __name__ == '__main__':
    print('Cessna %s:' % MODEL)
    total = 0
    for job in JOBS:
        total += build(*job)
    print('\nรวมจุดที่ต้นฉบับขาด %d จุด — ต้องสแกนหน้าที่ถูกตัดมาเพิ่มก่อนอนุมัติ' % total)
